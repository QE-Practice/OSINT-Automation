import os
import re
import time
from datetime import datetime
from selenium import webdriver
from selenium.webdriver import ActionChains
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.edge.service import Service
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Inches

# Path to the OSINT report Word template
word_template_path = r'C:\Users\gj115602\Downloads\OSINT_report_template.docx'

# Path to save the updated Word document
screenshot_folder = r'D:\Python Automation\OSINT Automation\Screenshots'

# Path to your Edge WebDriver and Wappalyzer extension
edge_driver_path = r'D:\Python Automation\OSINT Automation\msedgedriver.exe'

wappalyzer_extension_path = r'D:\Python Automation\OSINT Automation\Extensions\a02842b4-e272-4cff-937d-aee0e0ce3d5c.crx'

# Company name to replace in the document
new_company_name = "Facebook"
current_date = datetime.now().strftime("%B %d, %Y")  # Format: 'October 14, 2024'

# Function to replace text in paragraphs, runs, tables, headers, and footers
def replace_text_in_doc(doc, old_text, new_text):
    pattern = re.compile(re.escape(old_text), re.IGNORECASE)
    found = False

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            normalized_text = re.sub(r'\s+', ' ', run.text.strip())
            if pattern.search(normalized_text):
                run.text = pattern.sub(new_text, run.text)
                found = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        normalized_text = re.sub(r'\s+', ' ', run.text.strip())
                        if pattern.search(normalized_text):
                            run.text = pattern.sub(new_text, run.text)
                            found = True

    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                normalized_text = re.sub(r'\s+', ' ', run.text.strip())
                if pattern.search(normalized_text):
                    run.text = pattern.sub(new_text, run.text)
                    found = True

        footer = section.footer
        for paragraph in footer.paragraphs:
            for run in paragraph.runs:
                normalized_text = re.sub(r'\s+', ' ', run.text.strip())
                if pattern.search(normalized_text):
                    run.text = pattern.sub(new_text, run.text)
                    found = True

    if not found:
        print(f"No instances of '{old_text}' found in the document.")



# Function to visit a website, trigger Wappalyzer, and take a screenshot
def visit_website_and_take_wappalyzer_screenshot(url_to_scan, driver_path, extension_path, screenshot_folder):
    service = Service(driver_path)
    edge_options = webdriver.EdgeOptions()
    edge_options.add_argument("--inprivate")

    # Start Edge WebDriver
    driver = webdriver.Edge(service=service, options=edge_options)
    driver.get(url_to_scan)
    time.sleep(5)  # Wait for the page to load

    # Ensure the browser is in focus and activate the Wappalyzer extension using keyboard shortcut (Ctrl + Shift + W)
    driver.switch_to.active_element.click()
    actions = ActionChains(driver)
    actions.key_down(Keys.CONTROL).key_down(Keys.SHIFT).send_keys('W').key_up(Keys.SHIFT).key_up(Keys.CONTROL).perform()

    # Add a slight delay to ensure the extension activates
    time.sleep(3)

    # Take a screenshot
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    screenshot_filename = f"screenshot_wappalyzer_{timestamp}.png"
    screenshot_path = os.path.join(screenshot_folder, screenshot_filename)

    driver.save_screenshot(screenshot_path)
    driver.quit()

    return screenshot_path

# Function to go to SecurityHeaders.com, scan a site, and take a screenshot
def scan_security_headers_and_take_screenshot(url_to_scan, driver_path, screenshot_folder):
    service = Service(driver_path)
    edge_options = webdriver.EdgeOptions()
    driver = webdriver.Edge(service=service, options=edge_options)

    # Go to www.securityheaders.com
    driver.get("https://www.securityheaders.com")
    time.sleep(2)

    # Maximize the browser window
    driver.maximize_window()

    # Find the input box with ID 'q' and pass the target URL
    search_box = driver.find_element(By.ID, "q")
    search_box.send_keys(url_to_scan)
    time.sleep(1)

    # Find the Scan button and click it
    scan_button = driver.find_element(By.XPATH, "//input[@id='scan']")
    scan_button.click()
    time.sleep(2)

    # Scroll down to the Security Headers section
    security_headers_section = driver.find_element(By.XPATH, "//div[normalize-space()='Security Report Summary']")
    driver.execute_script("arguments[0].scrollIntoView();", security_headers_section)
    time.sleep(2)

    # Generate a dynamic filename for the screenshot using the current timestamp
    timestamp = time.strftime("%Y%m%d-%H%M%S")
    screenshot_filename = f"screenshot_{timestamp}.png"
    screenshot_path = os.path.join(screenshot_folder, screenshot_filename)

    # Capture screenshot and save it to the dynamically generated path
    driver.save_screenshot(screenshot_path)

    # Close the WebDriver after taking the screenshot
    driver.quit()

    return screenshot_path

# Function to insert the Wappalyzer and Security Headers screenshots into the Word document
def insert_screenshots_in_word(doc_template_path, wappalyzer_screenshot, sec_headers_screenshot, new_company_name, current_date, output_folder):
    # Open the existing Word document
    document = Document(doc_template_path)

    # Replace the placeholder company name and date in the document
    replace_text_in_doc(document, "Greenroads", new_company_name)
    replace_text_in_doc(document, "Date Placeholder", current_date)

    # Add a heading for the Wappalyzer section
    document.add_heading('Wappalyzer Analysis', level=1)

    # Add an explanatory paragraph before the Wappalyzer screenshot (optional)
    document.add_paragraph(
        "The following section provides a summary of the web technologies detected on the target website using the Wappalyzer tool."
    )

    # Insert the Wappalyzer screenshot
    document.add_picture(wappalyzer_screenshot, width=Inches(6))

    # Add a heading for the Security Headers section
    document.add_heading('Security Headers', level=1)

    # Add an explanatory paragraph before the Security Headers screenshot (optional)
    document.add_paragraph(
        "The following section provides a security header scan summary of the target website using SecurityHeaders.com."
    )

    # Insert the Security Headers screenshot
    document.add_picture(sec_headers_screenshot, width=Inches(6))

    # Generate a dynamic filename for the updated Word document
    timestamp = time.strftime("%Y%m%d-%H%M%S")
    updated_report_filename = f'{timestamp}_OSINT_report_{new_company_name}.docx'
    updated_report_path = os.path.join(output_folder, updated_report_filename)

    # Save the updated Word document
    document.save(updated_report_path)

    return updated_report_path

# Main execution flow: visit the website, trigger Wappalyzer, Security Headers scan, take screenshots, and insert them into the Word report
url_to_scan = "https://www.facebook.com"
wappalyzer_screenshot_path = visit_website_and_take_wappalyzer_screenshot(url_to_scan, edge_driver_path, wappalyzer_extension_path, screenshot_folder)
security_headers_screenshot_path = scan_security_headers_and_take_screenshot(url_to_scan, edge_driver_path, screenshot_folder)

# Insert the screenshots into the Word document
updated_report_path = insert_screenshots_in_word(
    word_template_path,
    wappalyzer_screenshot_path,
    security_headers_screenshot_path,
    new_company_name,
    current_date,
    screenshot_folder
)

print(f"Screenshots saved and report saved at {updated_report_path}")
